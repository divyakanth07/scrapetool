import requests
from bs4 import BeautifulSoup
from docx import Document

API_KEY = 'AIzaSyAppa-T4WCDc8WZ7bgp-jGDk1Db871Wi5E'
CSE_ID = '3770c62f79da54529'

def google_search(query, num_results):
    search_url = f"https://www.googleapis.com/customsearch/v1?q={query}&key={API_KEY}&cx={CSE_ID}&num={num_results}"
    response = requests.get(search_url)
    results = response.json()
    urls = [item['link'] for item in results.get('items', [])]
    return urls

def scrape_url(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        return soup.get_text()
    except requests.RequestException as e:
        return f"Failed to retrieve {url}: {e}"

def save_to_doc(url_data, filename):
    doc = Document()
    doc.add_heading('Scraped Data', 0)
    
    for url, data in url_data.items():
        doc.add_heading(url, level=1)
        doc.add_paragraph(data)
    
    doc.save(filename)

def main():
    query = input("Enter the keyword for Google search: ")
    num_results = 10
    filename = 'scraped_results.docx'
    
    urls = google_search(query, num_results)
    url_data = {url: scrape_url(url) for url in urls}
    
    save_to_doc(url_data, filename)
    print(f"Scraping complete. Results saved to {filename}")
    print_ascii_art()

def print_ascii_art():
    ascii_art = """
.------.------.------.------.     .------.------.     .------.------.------.------.------.------.------.------.------.------.
|M.--. |A.--. |D.--. |E.--. |.-.  |B.--. |Y.--. |.-.  |D.--. |I.--. |V.--. |Y.--. |A.--. |K.--. |A.--. |N.--. |T.--. |H.--. |
| (\/) | (\/) | :/\: | (\/) ((5)) | :(): | (\/) ((5)) | :/\: | (\/) | :(): | (\/) | (\/) | :/\: | (\/) | :(): | :/\: | :/\: |
| :\/: | :\/: | (__) | :\/: |'-.-.| ()() | :\/: |'-.-.| (__) | :\/: | ()() | :\/: | :\/: | :\/: | :\/: | ()() | (__) | (__) |
| '--'M| '--'A| '--'D| '--'E| ((1)| '--'B| '--'Y| ((1)| '--'D| '--'I| '--'V| '--'Y| '--'A| '--'K| '--'A| '--'N| '--'T| '--'H|
`------`------`------`------'  '-'`------`------'  '-'`------`------`------`------`------`------`------`------`------`------'"""
    print(ascii_art)

if __name__ == "__main__":
    main()
